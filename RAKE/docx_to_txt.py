import docx
from bs4 import BeautifulSoup
import requests
import os
import rake
import operator

class docx_to_txt:
    def __init__(self, filePath):
        self.name = filePath

    def parseDocx(self):
        doc = docx.Document(self.name)
        output_file_name = 'output.txt'
        fullText = []

        tables = doc.tables
        prev = ''
        definingCategories = ['Project Title','Project Manager', 'Project Start Date', 'Project End Date',
                              'Project Sponsor', 'Project Sponsor', 'Project Type', 'Function/Department'
                              'Operating Company/Division', 'Business Need', 'Project Scope', 'Deliverables'
                              'Risks & Issues','Assumptions','Key Activities','Financials','Milestones','Target Completion Date'
                              'Project Manager','Team Member','Sponsor','Corporate HR Manager','Operating Company HR','Operating Company President'
                              ,'Rank','Downloads','Shares']
        categoriesDict = {}
        saveInDict = False
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if(prev in definingCategories):
                            categoriesDict[prev] = paragraph.text
                        prev = paragraph.text

        return categoriesDict

class html_to_txt:
    def __init__(self, url):
        self.name = url

    def convertUrl(self):
        page = requests.get(self.name)
        soup = BeautifulSoup(page.content, 'html.parser')
        # print(soup.prettify())
        soup_string = str(soup.text)
        text_file = open("./output.txt","w")
        text_file.write(soup_string)
        text_file.close()


class rake_classify:
    def __init__(self, filePath):
        self.name = filePath

    def extractKeywords(self):
        with open(self.name, 'r') as myfile:
            text = myfile.read().replace('\n', '')
        # print "text is  ", text

        #using constraint where each keyword appears in text at least twice
        rake_object = rake.Rake("SmartStoplist.txt", 3, 3, 1)
        keywords = rake_object.run(text)
        print("keywords1 are ", keywords)

        #using constraint where each keyword appears in text at least three times
        rake_object = rake.Rake("SmartStoplist.txt", 3, 3, 2)
        keywords = rake_object.run(text)
        print("keywords2 are ", keywords)
